import os
import json
import yaml
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List, Union
from PIL import Image
import requests
from google import genai
from google.genai import types
import logging
from dataclasses import dataclass
from abc import ABC, abstractmethod

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ProcessedData:
    """Data structure for processed content"""
    content: Any
    metadata: Dict[str, Any]
    data_type: str
    source_file: str

class DataLoader(ABC):
    """Abstract base class for data loaders"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
    
    @abstractmethod
    def load(self, file_path: str) -> ProcessedData:
        """Load and process data from file"""
        pass
    
    @abstractmethod
    def supports(self, file_extension: str) -> bool:
        """Check if loader supports file extension"""
        pass

class ImageLoader(DataLoader):
    """Loader for image files"""
    
    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
    
    def load(self, file_path: str) -> ProcessedData:
        try:
            with open(file_path, 'rb') as f:
                image_bytes = f.read()
            
            image = Image.open(file_path)
            
            metadata = {
                'format': image.format,
                'size': image.size,
                'mode': image.mode,
                'file_size': len(image_bytes)
            }
            
            return ProcessedData(
                content=image_bytes,
                metadata=metadata,
                data_type='image',
                source_file=file_path
            )
        except Exception as e:
            logger.error(f"Error loading image {file_path}: {str(e)}")
            raise

class TextLoader(DataLoader):
    """Loader for text files"""
    
    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.txt', '.md', '.json', '.xml', '.html']
    
    def load(self, file_path: str) -> ProcessedData:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            metadata = {
                'file_size': len(text_content),
                'line_count': len(text_content.splitlines()),
                'encoding': 'utf-8'
            }
            
            return ProcessedData(
                content=text_content,
                metadata=metadata,
                data_type='text',
                source_file=file_path
            )
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            raise

class DocumentLoader(DataLoader):
    """Loader for document files (DOCX, PDF)"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self._install_dependencies()
    
    def _install_dependencies(self):
        try:
            import docx
            import PyPDF2
        except ImportError:
            logger.warning("Document processing libraries not found. Install: pip install python-docx pypdf2")
    
    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.docx', '.pdf']
    
    def load(self, file_path: str) -> ProcessedData:
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.docx':
                return self._load_docx(file_path)
            elif file_extension == '.pdf':
                return self._load_pdf(file_path)
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    def _load_docx(self, file_path: str) -> ProcessedData:
        from docx import Document
        
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            text_content.append(paragraph.text)
        
        full_text = '\n'.join(text_content)
        
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'file_size': os.path.getsize(file_path),
            'format': 'docx'
        }
        
        return ProcessedData(
            content=full_text,
            metadata=metadata,
            data_type='document',
            source_file=file_path
        )
    
    def _load_pdf(self, file_path: str) -> ProcessedData:
        import PyPDF2
        
        text_content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text_content.append(f"--- Page {page_num + 1} ---")
                text_content.append(page.extract_text())
        
        full_text = '\n'.join(text_content)
        
        metadata = {
            'page_count': len(pdf_reader.pages),
            'file_size': os.path.getsize(file_path),
            'format': 'pdf'
        }
        
        return ProcessedData(
            content=full_text,
            metadata=metadata,
            data_type='document',
            source_file=file_path
        )

class SpreadsheetLoader(DataLoader):
    """Loader for spreadsheet files (XLSX, CSV)"""
    
    def supports(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.xlsx', '.xls', '.csv']
    
    def load(self, file_path: str) -> ProcessedData:
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert to structured format
            data_dict = {
                'columns': df.columns.tolist(),
                'data': df.fillna('').to_dict('records'),
                'shape': df.shape
            }
            
            metadata = {
                'row_count': df.shape[0],
                'column_count': df.shape[1],
                'columns': df.columns.tolist(),
                'file_size': os.path.getsize(file_path),
                'format': file_extension[1:]
            }
            
            return ProcessedData(
                content=data_dict,
                metadata=metadata,
                data_type='spreadsheet',
                source_file=file_path
            )
        except Exception as e:
            logger.error(f"Error loading spreadsheet {file_path}: {str(e)}")
            raise

class DataLoaderFactory:
    """Factory class for creating appropriate data loaders"""
    
    def __init__(self, config: Dict[str, Any]):
        self.loaders = [
            ImageLoader(config),
            TextLoader(config),
            DocumentLoader(config),
            SpreadsheetLoader(config)
        ]
    
    def get_loader(self, file_path: str) -> DataLoader:
        """Get appropriate loader for file type"""
        file_extension = Path(file_path).suffix.lower()
        
        for loader in self.loaders:
            if loader.supports(file_extension):
                return loader
        
        raise ValueError(f"No loader available for file type: {file_extension}")
